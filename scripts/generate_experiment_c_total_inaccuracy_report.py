from __future__ import annotations

import hashlib
import json
import os
import platform
import shutil
import subprocess
from datetime import date, datetime, timezone
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", "/home/srinija/GPU_Datasets/.matplotlib-cache")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("data/final_report_experiment_c_v1")
FIG_DIR = OUT_DIR / "figures"
TABLE_DIR = OUT_DIR / "report_data_tables"
ANALYSIS_DIR = Path("data/analysis_experiment_c_final_v1")
FINAL_CSV = Path("data/grading_experiment_c_final_v1/final_scored_results.csv")
FINAL_JSONL = Path("data/grading_experiment_c_final_v1/final_scored_results.jsonl")
FINAL_DATASET_HASH = "6fdfaa035b5da2211e813353916902c871e783ecfa993615db672f62bcb8e327"
GRADER_HASH = "d9282a0ccc50daba3bfd232c058dfbab63a5c19a7323d585a7c2236b3a6c4ba8"
MODEL_ID = "meta-llama/Llama-3.2-3B-Instruct"
MODEL_REVISION = "0cb88a4f764b7a12671c53f0838cd831a0843b95"
PROMPT_VERSION = "llama_chat_v4"
PROMPT_HASH = "5d2869822989e19b"
CONTEXT_ORDER = ["4K", "8K", "16K", "32K", "64K"]
BANNED_TERMS = ["hallucinatory inaccuracy", "grounded inaccuracy"]


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def pct(x: float, decimals: int = 0) -> str:
    return f"{100 * x:.{decimals}f}%"


def fmt_p(p: float) -> str:
    return f"{p:.2e}" if p < 0.001 else f"{p:.4f}"


def load_data() -> dict:
    if sha256_file(FINAL_JSONL) != FINAL_DATASET_HASH:
        raise RuntimeError("final immutable dataset hash mismatch")
    df = pd.read_csv(FINAL_CSV)
    primary = pd.read_csv(ANALYSIS_DIR / "primary_results.csv")
    paired = pd.read_csv(ANALYSIS_DIR / "paired_tests.csv")
    qtype = pd.read_csv(ANALYSIS_DIR / "question_type_results.csv")
    domain = pd.read_csv(ANALYSIS_DIR / "domain_results.csv")
    latency = pd.read_csv(ANALYSIS_DIR / "latency_analysis.csv")
    models = json.loads((ANALYSIS_DIR / "mixed_model_results.json").read_text())
    if len(df) != 500 or df["instance_id"].nunique() != 500 or df["question_family_id"].nunique() != 100:
        raise RuntimeError("scored dataset integrity check failed")
    if any(df["context_length_label"].value_counts().get(c, 0) != 100 for c in CONTEXT_ORDER):
        raise RuntimeError("context-count integrity check failed")
    return {"df": df, "primary": primary, "paired": paired, "qtype": qtype, "domain": domain, "latency": latency, "models": models}


def prepare_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    for folder in [FIG_DIR, TABLE_DIR]:
        for path in folder.iterdir():
            if path.is_file():
                path.unlink()


def save_table(df: pd.DataFrame, name: str) -> None:
    df.to_csv(TABLE_DIR / f"{name}.csv", index=False)
    (TABLE_DIR / f"{name}.json").write_text(df.to_json(orient="records", indent=2), encoding="utf-8")


def build_tables(data: dict) -> dict[str, pd.DataFrame]:
    primary = data["primary"].set_index("context").loc[CONTEXT_ORDER].reset_index()
    latency = data["latency"].set_index("context").loc[CONTEXT_ORDER].reset_index()
    overall = pd.DataFrame(
        [
            {
                "Context": r["context"],
                "N": int(r["N"]),
                "Correct": pct(r["accuracy"]),
                "Inaccurate": pct(1 - r["accuracy"]),
                "Inaccuracy 95% CI": f"[{pct(1 - r['accuracy_ci_high'])}, {pct(1 - r['accuracy_ci_low'])}]",
                "Mean inference latency": f"{r['mean_latency']:.3f} s",
            }
            for _, r in primary.iterrows()
        ]
    )
    qdist = pd.DataFrame(
        [
            ["DIRECT_RETRIEVAL", 20, "Single-record factual retrieval."],
            ["RETRIEVAL_CALCULATION", 30, "Retrieve values and compute a deterministic answer."],
            ["TEMPORAL_VERSION", 11, "Identify the requested period, version, or submission."],
            ["ENTITY_UNIT_BINDING", 19, "Bind entity, unit, product, route, dosage form, or series attributes."],
            ["UNANSWERABLE", 20, "Return the insufficient-evidence answer when target evidence is absent."],
        ],
        columns=["Question type", "Families", "Purpose"],
    )
    domains = pd.DataFrame(
        [
            ["SEC", "Company financial facts and filings."],
            ["FDA / Drugs@FDA", "Drug applications, submissions, products, strengths, dosage forms, and routes."],
            ["ClinicalTrials.gov", "Trial metadata, dates, enrollment, arms, status, and unavailable fields."],
            ["FRED", "Economic time series with unit, frequency, and series attributes."],
        ],
        columns=["Domain", "Role"],
    )
    paired = data["paired"]
    paired = paired[paired["outcome"] == "answer_correct_int"].copy()
    paired_rows = []
    for _, r in paired.iterrows():
        diff = -r["absolute_paired_rate_difference"]
        paired_rows.append(
            {
                "Comparison": r["comparison"].replace("_", " "),
                "Inaccuracy difference": pct(diff, 1),
                "Discordant pairs": f"{int(r['discordant_4k_false_comparison_true'])}/{int(r['discordant_4k_true_comparison_false'])}",
                "Raw p-value": fmt_p(r["raw_p_value"]),
                "Holm-adjusted p-value": fmt_p(r["holm_adjusted_p_value"]),
                "Result": "p < 0.05 after Holm" if r["holm_adjusted_p_value"] < 0.05 else "not significant after Holm",
            }
        )
    paired_table = pd.DataFrame(paired_rows)
    qtype = subgroup_table(data["df"], "question_type")
    domain = subgroup_table(data["df"], "domain")
    latency_table = latency[["context", "mean_latency", "median_latency", "p95_latency", "total_latency", "mean_input_tokens", "mean_generated_tokens"]].copy()
    latency_table.columns = ["Context", "Mean latency (s)", "Median latency (s)", "P95 latency (s)", "Total latency (s)", "Mean input tokens", "Mean output tokens"]
    tables = {
        "overall_results_total_inaccuracy": overall,
        "question_type_distribution": qdist,
        "domain_summary": domains,
        "paired_inaccuracy_tests": paired_table,
        "question_type_inaccuracy": qtype,
        "domain_inaccuracy": domain,
        "latency_summary": latency_table,
    }
    for name, table in tables.items():
        save_table(table, name)
    return tables


def subgroup_table(df: pd.DataFrame, group_col: str) -> pd.DataFrame:
    rows = []
    for (group, context), g in df.groupby([group_col, "context_length_label"], sort=True):
        acc = float(g["final_answer_correct"].mean())
        rows.append({"Group": group, "Context": context, "N": len(g), "Correct": pct(acc, 1), "Inaccurate": pct(1 - acc, 1)})
    return pd.DataFrame(rows)


def make_figures(data: dict) -> list[dict[str, str]]:
    primary = data["primary"].set_index("context").loc[CONTEXT_ORDER].reset_index()
    latency = data["latency"].set_index("context").loc[CONTEXT_ORDER].reset_index()
    df = data["df"]
    models = data["models"]
    acc = models["primary_accuracy"]
    inacc_or = 1 / acc["odds_ratio_per_context_doubling"]
    inacc_ci = [1 / acc["odds_ratio_95_ci"][1], 1 / acc["odds_ratio_95_ci"][0]]
    x = np.arange(len(CONTEXT_ORDER))
    figs: list[dict[str, str]] = []

    def save(name: str, title: str, caption: str) -> None:
        png = FIG_DIR / f"{name}.png"
        pdf = FIG_DIR / f"{name}.pdf"
        plt.tight_layout()
        plt.savefig(png, dpi=300)
        plt.savefig(pdf)
        plt.close()
        figs.append({"name": name, "title": title, "caption": caption, "png": str(png), "pdf": str(pdf)})

    correct = primary["accuracy"].to_numpy()
    inaccurate = 1 - correct
    plt.figure(figsize=(8, 4.8))
    plt.bar(x, correct, label="Correct", color="#2b8cbe")
    plt.bar(x, inaccurate, bottom=correct, label="Inaccurate", color="#d95f02")
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylim(0, 1)
    plt.ylabel("Share of responses")
    plt.xlabel("Context length")
    plt.title("Factual Reliability vs Context Length")
    plt.legend(loc="upper center", ncol=2, bbox_to_anchor=(0.5, -0.12))
    save("figure_1_factual_reliability_total_inaccuracy", "Figure 1. Factual reliability vs context length.", "Correct and inaccurate responses sum to 100% at each context length.")

    inacc_low = 1 - primary["accuracy_ci_high"]
    inacc_high = 1 - primary["accuracy_ci_low"]
    plt.figure(figsize=(7.5, 4.7))
    plt.errorbar(x, inaccurate, yerr=[inaccurate - inacc_low, inacc_high - inaccurate], marker="o", capsize=4, color="#7b3294")
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylim(0, 1)
    plt.ylabel("Inaccuracy rate")
    plt.xlabel("Context length")
    plt.title("Inaccuracy vs Context Length")
    save("figure_2_inaccuracy", "Figure 2. Inaccuracy vs context length.", f"GEE OR per context doubling = {inacc_or:.3f}, 95% CI [{inacc_ci[0]:.3f}, {inacc_ci[1]:.3f}], p = {fmt_p(acc['p_value_context_log2'])}.")

    plt.figure(figsize=(7.5, 4.7))
    plt.plot(x, latency["mean_latency"], marker="o", color="#1b7837")
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylabel("Mean synchronized inference latency (s)")
    plt.xlabel("Context length")
    plt.title("Inference Latency vs Context Length")
    save("figure_3_latency", "Figure 3. Inference latency vs context length.", "Mean synchronized generation latency rises as input length increases.")

    for group_col, fig_name, title in [
        ("question_type", "figure_4_inaccuracy_by_question_type", "Inaccuracy by Question Type and Context"),
        ("domain", "figure_5_inaccuracy_by_domain", "Inaccuracy by Domain and Context"),
    ]:
        rows = []
        for (group, context), g in df.groupby([group_col, "context_length_label"], sort=True):
            rows.append({group_col: group, "context": context, "inaccuracy": 1 - g["final_answer_correct"].mean()})
        plot_df = pd.DataFrame(rows)
        plt.figure(figsize=(8.8, 5))
        for group, g in plot_df.groupby(group_col):
            g = g.set_index("context").loc[CONTEXT_ORDER].reset_index()
            plt.plot(x, g["inaccuracy"], marker="o", label=group)
        plt.xticks(x, CONTEXT_ORDER)
        plt.ylim(0, 1)
        plt.ylabel("Inaccuracy rate")
        plt.xlabel("Context length")
        plt.title(title)
        plt.legend(fontsize=8, ncol=2)
        save(fig_name, title.replace(" and ", " & ") + ".", "Subgroup results are exploratory because cells are smaller than the primary repeated-measures analysis.")
    return figs


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        s = doc.styles[style_name]
        s.font.name = "Aptos"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(31, 78, 121)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_table(doc: Document, df: pd.DataFrame, title: str, max_rows: int | None = None) -> None:
    doc.add_paragraph(title, style="Heading 3")
    show = df if max_rows is None else df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Light Shading Accent 1"
    for i, col in enumerate(show.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(show.columns):
            cells[i].text = str(row[col])
    doc.add_paragraph()


def add_figure(doc: Document, fig: dict[str, str], width: float = 6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(fig["title"]).bold = True
    doc.add_picture(fig["png"], width=Inches(width))
    cap = doc.add_paragraph(fig["caption"])
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(data: dict, tables: dict[str, pd.DataFrame], figures: list[dict[str, str]]) -> Path:
    primary = data["primary"].set_index("context").loc[CONTEXT_ORDER].reset_index()
    models = data["models"]
    acc = models["primary_accuracy"]
    inacc_or = 1 / acc["odds_ratio_per_context_doubling"]
    inacc_ci = [1 / acc["odds_ratio_95_ci"][1], 1 / acc["odds_ratio_95_ci"][0]]
    doc = Document()
    style_doc(doc)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Long-Context Factual Reliability in Llama 3.2 3B Instruct")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(31, 78, 121)
    for text in [
        "Experiment C: Total Inaccuracy Report",
        f"Model: {MODEL_ID}",
        "Dataset: 100 question families, 500 repeated observations",
        f"Report date: {date.today().isoformat()}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    for s in [
        "1. Executive Summary / Abstract",
        "2. Research Question and Motivation",
        "3. Dataset and Source Domains",
        "4. Question Dataset",
        "5. Experimental Configuration",
        "6. Grading Method",
        "7. Overall Results",
        "8. Statistical Analysis of Overall Inaccuracy",
        "9. Paired Context Comparisons",
        "10. Question-Type Results",
        "11. Domain Results",
        "12. Inference-Time Analysis",
        "13. Interpretation, Limitations, and Conclusion",
        "14. Methodological Provenance",
        "Appendix A. Detailed Tables",
    ]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_page_break()

    doc.add_heading("1. Executive Summary / Abstract", level=1)
    doc.add_paragraph(
        f"This report summarizes Experiment C, a completed local inference experiment evaluating factual reliability for {MODEL_ID}. "
        "The benchmark uses primary-source records from SEC filings, Drugs@FDA, ClinicalTrials.gov, and FRED to create 100 repeated "
        "question families evaluated at 4K, 8K, 16K, 32K, and 64K context lengths, for 500 total observations."
    )
    doc.add_paragraph(
        f"The main result is a strong increase in total inaccuracy as context length increases. Inaccuracy rose from 43% at 4K to 73% "
        f"at 64K. In a GEE logistic model clustered by question family, each doubling of context length was associated with "
        f"{inacc_or:.2f}x the odds of an inaccurate answer (95% CI [{inacc_ci[0]:.2f}, {inacc_ci[1]:.2f}], "
        f"p = {fmt_p(acc['p_value_context_log2'])}). Mean inference latency increased from 0.317 s to 7.824 s."
    )

    doc.add_heading("2. Research Question and Motivation", level=1)
    doc.add_paragraph(
        "The primary research question is how increasing context length affects factual reliability when a model is given legitimate "
        "but competing same-domain records. The report uses total inaccuracy as the main reliability metric: a response is either "
        "Correct or Inaccurate."
    )

    doc.add_heading("3. Dataset and Source Domains", level=1)
    doc.add_paragraph(
        "The benchmark converts authoritative primary-source records into controlled factual questions. Each question family has a "
        "deterministic gold answer, target evidence, structured target conditions, and same-domain distractors. Contexts grow by "
        "adding plausible competing information while preserving the same question, gold answer, and target evidence."
    )
    add_table(doc, tables["domain_summary"], "Table 1. Source domains.")

    doc.add_heading("4. Question Dataset", level=1)
    doc.add_paragraph(
        "The final dataset contains 100 unique question families: 80 answerable and 20 unanswerable. Each family was evaluated at five "
        "context lengths, yielding 100 x 5 = 500 repeated observations. Because the same families appear at all context lengths, "
        "statistical tests treat observations as clustered by family."
    )
    add_table(doc, tables["question_type_distribution"], "Table 2. Question-type distribution.")

    doc.add_heading("5. Experimental Configuration", level=1)
    config = pd.DataFrame(
        [
            ["Model", MODEL_ID],
            ["Model revision", MODEL_REVISION],
            ["Hardware", "NVIDIA GeForce RTX 4090"],
            ["Precision/cache", "BF16, standard DynamicCache"],
            ["Decoding", "deterministic greedy; do_sample=False; num_beams=1; max_new_tokens=128"],
            ["Quantization/offloading", "none"],
            ["Prompt version/hash", f"{PROMPT_VERSION} / {PROMPT_HASH}"],
            ["Output contract", "ANSWER: <answer>"],
        ],
        columns=["Setting", "Value"],
    )
    add_table(doc, config, "Table 3. Inference configuration.")
    doc.add_paragraph("Experiment C achieved 500/500 successful inference instances, zero format failures, and zero outputs hitting the 128-token cap.")

    doc.add_heading("6. Grading Method", level=1)
    doc.add_paragraph(
        f"Responses were graded by a frozen deterministic grader (SHA-256 {GRADER_HASH}). Of 500 responses, 499 were resolved by the "
        "deterministic grader and one was resolved by documented human manual adjudication. The report presents total inaccuracy as "
        "the reliability outcome: an answer is Correct if it matches the deterministic gold answer after normalization, otherwise "
        "it is Inaccurate."
    )

    doc.add_heading("7. Overall Results", level=1)
    add_table(doc, tables["overall_results_total_inaccuracy"], "Table 4. Central factual-reliability results.")
    add_figure(doc, figures[0])
    add_figure(doc, figures[1])

    doc.add_heading("8. Statistical Analysis of Overall Inaccuracy", level=1)
    doc.add_paragraph(
        "The repeated-measures model was fit as a GEE logistic regression clustered by question_family_id with context_log2 coded as "
        "4K = 0, 8K = 1, 16K = 2, 32K = 3, and 64K = 4."
    )
    doc.add_paragraph(
        f"Accuracy OR per context doubling was {acc['odds_ratio_per_context_doubling']:.3f} "
        f"(95% CI [{acc['odds_ratio_95_ci'][0]:.3f}, {acc['odds_ratio_95_ci'][1]:.3f}], "
        f"p = {fmt_p(acc['p_value_context_log2'])}). Equivalently, inaccuracy OR per context doubling was {inacc_or:.3f} "
        f"(95% CI [{inacc_ci[0]:.3f}, {inacc_ci[1]:.3f}], p = {fmt_p(acc['p_value_context_log2'])}). "
        "This is an odds increase, not a percentage-point increase."
    )

    doc.add_heading("9. Paired Context Comparisons", level=1)
    doc.add_paragraph("Paired McNemar tests compared each longer context to 4K within the same question families. Holm correction was applied across the four comparisons.")
    add_table(doc, tables["paired_inaccuracy_tests"], "Table 5. Paired inaccuracy comparisons.")

    doc.add_heading("10. Question-Type Results", level=1)
    doc.add_paragraph("Question-type analyses are exploratory because subgroup cell sizes are smaller than the primary repeated-measures analysis.")
    add_figure(doc, figures[3])
    add_table(doc, tables["question_type_inaccuracy"], "Table 6. Inaccuracy by question type and context.", max_rows=25)

    doc.add_heading("11. Domain Results", level=1)
    doc.add_paragraph("Domain analyses are exploratory because each domain contributes 25 families.")
    add_figure(doc, figures[4])
    add_table(doc, tables["domain_inaccuracy"], "Table 7. Inaccuracy by domain and context.")

    doc.add_heading("12. Inference-Time Analysis", level=1)
    add_figure(doc, figures[2])
    add_table(doc, tables["latency_summary"], "Table 8. Inference latency summary.")
    doc.add_paragraph("Latency grows sharply with context length. This analysis is descriptive and should not be interpreted as causal evidence about reliability.")

    doc.add_heading("13. Interpretation, Limitations, and Conclusion", level=1)
    doc.add_paragraph(
        "Longer contexts substantially reduce factual reliability in Llama 3.2 3B under this benchmark. Overall inaccuracy increased "
        "from 43% at 4K to 73% at 64K, and the odds of an inaccurate answer increased by approximately 35% per context doubling."
    )
    for item in [
        "The experiment includes 100 question families and 500 observations.",
        "Only one model was evaluated: Llama 3.2 3B Instruct.",
        "Experiment C covers contexts through 64K.",
        "The benchmark uses four factual/structured data domains and intentionally strong same-domain distractors.",
        "Experiment C uses an answer-only format and therefore has no evidence-selection metric.",
        "Grading is deterministic with one documented manual adjudication.",
        "Subgroup analyses have smaller samples and are exploratory.",
        "GEE was used instead of mixed-effects logistic regression for reliable repeated-measures fitting.",
        "Results may not generalize to larger, proprietary, or differently prompted models.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("14. Methodological Provenance", level=1)
    provenance = pd.DataFrame(
        [
            ["Final scored dataset hash", FINAL_DATASET_HASH],
            ["Frozen grader hash", GRADER_HASH],
            ["Analysis directory", str(ANALYSIS_DIR)],
            ["Statistical method", "GEE logistic regression clustered by question_family_id; paired McNemar tests"],
            ["Bootstrap CIs", "Family-clustered bootstrap; seed 20260810; 5000 replicates"],
            ["No reruns", "Inference and grading were not rerun for this report."],
        ],
        columns=["Item", "Value"],
    )
    add_table(doc, provenance, "Table 9. Provenance.")

    doc.add_page_break()
    doc.add_heading("Appendix A. Detailed Tables", level=1)
    add_table(doc, tables["question_type_inaccuracy"], "Appendix Table A1. Full question-type results.")
    add_table(doc, tables["domain_inaccuracy"], "Appendix Table A2. Full domain results.")

    out = OUT_DIR / "final_research_report.docx"
    doc.save(out)
    return out


def build_markdown(data: dict, tables: dict[str, pd.DataFrame], figures: list[dict[str, str]]) -> Path:
    acc = data["models"]["primary_accuracy"]
    inacc_or = 1 / acc["odds_ratio_per_context_doubling"]
    inacc_ci = [1 / acc["odds_ratio_95_ci"][1], 1 / acc["odds_ratio_95_ci"][0]]
    lines = [
        "# Long-Context Factual Reliability in Llama 3.2 3B Instruct",
        "",
        "## Executive Summary",
        "",
        f"Total inaccuracy increased from 43% at 4K to 73% at 64K. GEE-estimated inaccuracy OR per context doubling was {inacc_or:.3f} (95% CI [{inacc_ci[0]:.3f}, {inacc_ci[1]:.3f}], p={fmt_p(acc['p_value_context_log2'])}).",
        "",
        "## Central Results",
        "",
        tables["overall_results_total_inaccuracy"].to_markdown(index=False),
        "",
        "## Figures",
        "",
    ]
    for i, fig in enumerate(figures, 1):
        lines += [f"### Figure {i}. {fig['title']}", "", f"![{fig['title']}]({Path(fig['png']).relative_to(OUT_DIR)})", "", fig["caption"], ""]
    lines += [
        "## Provenance",
        "",
        f"- Final dataset hash: `{FINAL_DATASET_HASH}`",
        f"- Frozen grader hash: `{GRADER_HASH}`",
    ]
    out = OUT_DIR / "final_research_report.md"
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def export_pdf(docx_path: Path) -> Path | None:
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        return None
    proc = subprocess.run(
        [
            soffice,
            "--headless",
            "-env:UserInstallation=file:///tmp/libreoffice-profile",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT_DIR.resolve()),
            str(docx_path.resolve()),
        ],
        text=True,
        capture_output=True,
        timeout=120,
    )
    pdf = OUT_DIR / "final_research_report.pdf"
    if proc.returncode != 0 or not pdf.exists():
        (OUT_DIR / "pdf_export_error.txt").write_text(proc.stdout + "\n" + proc.stderr, encoding="utf-8")
        return None
    return pdf


def verify_docx(path: Path, expected_figures: int) -> dict:
    d = Document(path)
    text = "\n".join(p.text for p in d.paragraphs)
    lower = text.casefold()
    required = ["Executive Summary / Abstract", "Research Question and Motivation", "Overall Results", "Statistical Analysis of Overall Inaccuracy", "Inference-Time Analysis", "Methodological Provenance"]
    return {
        "docx_readable": True,
        "paragraph_count": len(d.paragraphs),
        "table_count": len(d.tables),
        "embedded_figure_count": len(d.inline_shapes),
        "missing_required_sections": [r for r in required if r not in text],
        "figures_embedded": len(d.inline_shapes) >= expected_figures,
        "banned_terms_present": [term for term in BANNED_TERMS if term in lower],
        "total_inaccuracy_language_present": "total inaccuracy" in lower,
    }


def versions() -> dict[str, str]:
    import docx
    import matplotlib
    import numpy
    import pandas

    return {
        "python": platform.python_version(),
        "platform": platform.platform(),
        "python_docx": docx.__version__,
        "matplotlib": matplotlib.__version__,
        "numpy": numpy.__version__,
        "pandas": pandas.__version__,
    }


def main() -> int:
    prepare_dirs()
    data = load_data()
    tables = build_tables(data)
    figures = make_figures(data)
    docx = build_docx(data, tables, figures)
    md = build_markdown(data, tables, figures)
    pdf = export_pdf(docx)
    check = verify_docx(docx, len(figures))
    manifest = {
        "created_at": datetime.now().isoformat(),
        "source_dataset_hash": sha256_file(FINAL_JSONL),
        "expected_source_dataset_hash": FINAL_DATASET_HASH,
        "grader_hash": GRADER_HASH,
        "analysis_directory": str(ANALYSIS_DIR),
        "report_generation_script": "scripts/generate_experiment_c_total_inaccuracy_report.py",
        "report_generation_script_sha256": sha256_file(Path("scripts/generate_experiment_c_total_inaccuracy_report.py")),
        "package_versions": versions(),
        "figure_files": figures,
        "output_files": {
            "docx": str(docx),
            "markdown": str(md),
            "pdf": str(pdf) if pdf else None,
            "figures": str(FIG_DIR),
            "tables": str(TABLE_DIR),
        },
        "quality_check": check,
        "terminology_policy": "Total inaccuracy only; banned terms removed.",
        "inference_rerun": False,
        "grading_rerun": False,
        "labels_modified": False,
    }
    (OUT_DIR / "report_manifest.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False, sort_keys=True) + "\n", encoding="utf-8")
    print(json.dumps({"docx": str(docx), "markdown": str(md), "pdf": str(pdf) if pdf else None, "figures": len(figures), "quality_check": check, "manifest": str(OUT_DIR / "report_manifest.json")}, indent=2), flush=True)
    return 0 if not check["missing_required_sections"] and not check["banned_terms_present"] and check["figures_embedded"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
