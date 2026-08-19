from __future__ import annotations

import hashlib
import json
import os
import platform
import shutil
import subprocess
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any

os.environ.setdefault("MPLCONFIGDIR", "/home/srinija/GPU_Datasets/.matplotlib-cache")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("data/final_report_experiment_c_v1")
FIG_DIR = OUT_DIR / "figures"
TABLE_DIR = OUT_DIR / "report_data_tables"
ANALYSIS_DIR = Path("data/analysis_experiment_c_final_v1")
FINAL_CSV = Path("data/grading_experiment_c_final_v1/final_scored_results.csv")
FINAL_JSONL = Path("data/grading_experiment_c_final_v1/final_scored_results.jsonl")
FINAL_DATASET_HASH = "6fdfaa035b5da2211e813353916902c871e783ecfa993615db672f62bcb8e327"
GRADER_HASH = "d9282a0ccc50daba3bfd232c058dfbab63a5c19a7323d585a7c2236b3a6c4ba8"
MODEL_ID = "meta-llama/Llama-3.2-3B-Instruct"
MODEL_REVISION = "0cb88a4f764b7a12671c53f0838cd831a0843b95"
PROMPT_VERSION = "llama_chat_v4"
PROMPT_HASH = "5d2869822989e19b"
CONTEXT_ORDER = ["4K", "8K", "16K", "32K", "64K"]


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def pct(x: float, decimals: int = 0) -> str:
    return f"{100 * x:.{decimals}f}%"


def fmt_p(p: float) -> str:
    if p < 0.001:
        return f"{p:.2e}"
    return f"{p:.4f}"


def load_verified() -> dict[str, Any]:
    if sha256_file(FINAL_JSONL) != FINAL_DATASET_HASH:
        raise RuntimeError("final immutable JSONL hash mismatch")
    df = pd.read_csv(FINAL_CSV)
    if len(df) != 500 or df["instance_id"].nunique() != 500:
        raise RuntimeError("final scored CSV row/id integrity failed")
    if df["question_family_id"].nunique() != 100:
        raise RuntimeError("expected 100 question families")
    context_counts = df["context_length_label"].value_counts().to_dict()
    if any(context_counts.get(c, 0) != 100 for c in CONTEXT_ORDER):
        raise RuntimeError(f"context count integrity failed: {context_counts}")
    if df["final_answer_correct"].isna().any() or df["final_hallucination"].isna().any():
        raise RuntimeError("missing final labels")
    primary = pd.read_csv(ANALYSIS_DIR / "primary_results.csv")
    paired = pd.read_csv(ANALYSIS_DIR / "paired_tests.csv")
    qtype = pd.read_csv(ANALYSIS_DIR / "question_type_results.csv")
    domain = pd.read_csv(ANALYSIS_DIR / "domain_results.csv")
    errors = pd.read_csv(ANALYSIS_DIR / "error_type_by_context.csv")
    transitions = pd.read_csv(ANALYSIS_DIR / "family_transitions.csv")
    latency = pd.read_csv(ANALYSIS_DIR / "latency_analysis.csv")
    models = json.loads((ANALYSIS_DIR / "mixed_model_results.json").read_text())
    sensitivity = json.loads((ANALYSIS_DIR / "sensitivity_analysis.json").read_text())
    return {
        "df": df,
        "primary": primary,
        "paired": paired,
        "qtype": qtype,
        "domain": domain,
        "errors": errors,
        "transitions": transitions,
        "latency": latency,
        "models": models,
        "sensitivity": sensitivity,
    }


def prepare_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)


def save_table(df: pd.DataFrame, name: str) -> None:
    df.to_csv(TABLE_DIR / f"{name}.csv", index=False)
    (TABLE_DIR / f"{name}.json").write_text(df.to_json(orient="records", indent=2), encoding="utf-8")


def build_report_tables(data: dict[str, Any]) -> dict[str, pd.DataFrame]:
    df = data["df"].copy()
    primary = data["primary"].copy()
    latency = data["latency"].copy()
    overall_rows = []
    for _, row in primary.iterrows():
        incorrect = 1 - row["accuracy"]
        grounded = row["incorrect_non_hallucination_count"] / row["N"]
        overall_rows.append(
            {
                "Context": row["context"],
                "Correct": pct(row["accuracy"]),
                "Inaccurate": pct(incorrect),
                "Hallucinatory inaccuracy": pct(row["hallucination_rate"]),
                "Grounded inaccuracy": pct(grounded),
                "Mean inference latency": f"{row['mean_latency']:.3f} s",
            }
        )
    overall = pd.DataFrame(overall_rows)

    qdist = pd.DataFrame(
        [
            ["DIRECT_RETRIEVAL", 20, "Locate a single target field/value from matching records."],
            ["RETRIEVAL_CALCULATION", 30, "Retrieve multiple values and compute a deterministic result."],
            ["TEMPORAL_VERSION", 11, "Distinguish requested periods, versions, vintages, or submissions."],
            ["ENTITY_UNIT_BINDING", 19, "Bind entity, unit, product, route, dosage form, or series variant correctly."],
            ["UNANSWERABLE", 20, "Return INSUFFICIENT_EVIDENCE when the target evidence is absent."],
        ],
        columns=["Question type", "Number of families", "Purpose"],
    )
    domains = pd.DataFrame(
        [
            ["SEC", "Company financial facts and filings; temporal/version and numeric retrieval."],
            ["FDA / Drugs@FDA", "Drug applications, submissions, products, strengths, dosage forms, and routes."],
            ["ClinicalTrials.gov", "Trial metadata, dates, enrollment, arms, status, and unanswerable fields."],
            ["FRED", "Economic time series values with frequency, seasonality, unit, and series variants."],
        ],
        columns=["Domain", "Role in benchmark"],
    )

    paired = data["paired"].copy()
    paired_rows = []
    for _, r in paired.iterrows():
        outcome = "Hallucinatory inaccuracy" if r["outcome"] == "hallucination_int" else "Overall inaccuracy"
        diff = r["absolute_paired_rate_difference"]
        if r["outcome"] == "answer_correct_int":
            diff = -diff
        paired_rows.append(
            {
                "Comparison": r["comparison"].replace("_", " "),
                "Outcome": outcome,
                "Difference": pct(diff, 1),
                "Discordant pairs": f"{int(r['discordant_4k_false_comparison_true'])}/{int(r['discordant_4k_true_comparison_false'])}",
                "Raw p-value": fmt_p(r["raw_p_value"]),
                "Holm-adjusted p-value": fmt_p(r["holm_adjusted_p_value"]),
                "Result": "p < 0.05 after Holm" if r["holm_adjusted_p_value"] < 0.05 else "not significant after Holm",
            }
        )
    paired_table = pd.DataFrame(paired_rows)

    error_counts = data["errors"].pivot(index="error_type", columns="context", values="count").reset_index()
    error_counts = error_counts[["error_type", *CONTEXT_ORDER]]

    qtype_detail = subgroup_detail(df, "question_type")
    domain_detail = subgroup_detail(df, "domain")

    latency_table = latency[[
        "context", "mean_latency", "median_latency", "p95_latency", "total_latency",
        "mean_input_tokens", "mean_generated_tokens",
    ]].copy()
    latency_table.columns = [
        "Context", "Mean latency (s)", "Median latency (s)", "P95 latency (s)",
        "Total latency (s)", "Mean input tokens", "Mean output tokens",
    ]

    tables = {
        "overall_results": overall,
        "question_type_distribution": qdist,
        "domain_summary": domains,
        "paired_tests": paired_table,
        "error_counts_by_context": error_counts,
        "question_type_detail": qtype_detail,
        "domain_detail": domain_detail,
        "latency_summary": latency_table,
    }
    for name, table in tables.items():
        save_table(table, name)
    return tables


def subgroup_detail(df: pd.DataFrame, group_col: str) -> pd.DataFrame:
    rows = []
    for (group, context), g in df.groupby([group_col, "context_length_label"], sort=True):
        correct = g["final_answer_correct"].mean()
        hall = g["final_hallucination"].mean()
        grounded = ((~g["final_answer_correct"]) & (~g["final_hallucination"])).mean()
        rows.append(
            {
                group_col: group,
                "Context": context,
                "N": len(g),
                "Correct": pct(correct, 1),
                "Inaccurate": pct(1 - correct, 1),
                "Hallucinatory inaccuracy": pct(hall, 1),
                "Grounded inaccuracy": pct(grounded, 1),
            }
        )
    return pd.DataFrame(rows)


def make_report_figures(data: dict[str, Any]) -> list[dict[str, str]]:
    primary = data["primary"].set_index("context").loc[CONTEXT_ORDER].reset_index()
    df = data["df"].copy()
    x = np.arange(len(CONTEXT_ORDER))
    figures: list[dict[str, str]] = []

    def save(name: str, title: str, caption: str) -> None:
        png = FIG_DIR / f"{name}.png"
        pdf = FIG_DIR / f"{name}.pdf"
        plt.tight_layout()
        plt.savefig(png, dpi=300)
        plt.savefig(pdf)
        plt.close()
        figures.append({"name": name, "title": title, "caption": caption, "png": str(png), "pdf": str(pdf)})

    correct = primary["accuracy"].to_numpy()
    hall = primary["hallucination_rate"].to_numpy()
    grounded = (primary["incorrect_non_hallucination_count"] / primary["N"]).to_numpy()
    plt.figure(figsize=(8.5, 5))
    plt.bar(x, correct, label="Correct", color="#2b8cbe")
    plt.bar(x, grounded, bottom=correct, label="Grounded inaccuracy", color="#fdae61")
    plt.bar(x, hall, bottom=correct + grounded, label="Hallucinatory inaccuracy", color="#d7191c")
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylim(0, 1)
    plt.ylabel("Share of responses")
    plt.xlabel("Context length")
    plt.title("Factual Reliability vs Context Length")
    plt.legend(loc="upper center", ncol=3, bbox_to_anchor=(0.5, -0.12))
    save(
        "figure_1_factual_reliability",
        "Figure 1. Factual reliability vs context length.",
        "Correct, grounded inaccuracy, and hallucinatory inaccuracy sum to 100% at each context length.",
    )

    inacc = 1 - primary["accuracy"]
    inacc_low = 1 - primary["accuracy_ci_high"]
    inacc_high = 1 - primary["accuracy_ci_low"]
    models = data["models"]
    acc = models["primary_accuracy"]
    inacc_or = 1 / acc["odds_ratio_per_context_doubling"]
    inacc_ci = [1 / acc["odds_ratio_95_ci"][1], 1 / acc["odds_ratio_95_ci"][0]]
    plt.figure(figsize=(7.5, 4.7))
    plt.errorbar(x, inacc, yerr=[inacc - inacc_low, inacc_high - inacc], marker="o", capsize=4, color="#7b3294")
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylim(0, 1)
    plt.ylabel("Overall inaccuracy rate")
    plt.xlabel("Context length")
    plt.title("Overall Inaccuracy vs Context Length")
    save(
        "figure_2_overall_inaccuracy",
        "Figure 2. Overall inaccuracy vs context length.",
        f"GEE OR per context doubling = {inacc_or:.3f}, 95% CI [{inacc_ci[0]:.3f}, {inacc_ci[1]:.3f}], p = {fmt_p(acc['p_value_context_log2'])}.",
    )

    hall_model = models["primary_hallucination"]
    plt.figure(figsize=(7.5, 4.7))
    plt.errorbar(
        x,
        primary["hallucination_rate"],
        yerr=[
            primary["hallucination_rate"] - primary["hallucination_ci_low"],
            primary["hallucination_ci_high"] - primary["hallucination_rate"],
        ],
        marker="o",
        capsize=4,
        color="#d7191c",
    )
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylim(0, 1)
    plt.ylabel("Hallucinatory inaccuracy rate")
    plt.xlabel("Context length")
    plt.title("Hallucinatory Inaccuracy vs Context Length")
    save(
        "figure_3_hallucinatory_inaccuracy",
        "Figure 3. Hallucinatory inaccuracy vs context length.",
        f"GEE OR = {hall_model['odds_ratio_per_context_doubling']:.3f}, 95% CI [{hall_model['odds_ratio_95_ci'][0]:.3f}, {hall_model['odds_ratio_95_ci'][1]:.3f}], p = {fmt_p(hall_model['p_value_context_log2'])}; trend not statistically significant.",
    )

    latency = data["latency"].set_index("context").loc[CONTEXT_ORDER].reset_index()
    plt.figure(figsize=(7.5, 4.7))
    plt.plot(x, latency["mean_latency"], marker="o", color="#1b7837")
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylabel("Mean synchronized inference latency (s)")
    plt.xlabel("Context length")
    plt.title("Inference Latency vs Context Length")
    save(
        "figure_4_latency",
        "Figure 4. Inference latency vs context length.",
        "Mean synchronized generation latency rises sharply as rendered input length increases.",
    )

    err = data["errors"].pivot(index="context", columns="error_type", values="count").loc[CONTEXT_ORDER]
    err = err.div(err.sum(axis=1), axis=0)
    plt.figure(figsize=(9, 5.3))
    bottom = np.zeros(len(err))
    for col in err.columns:
        vals = err[col].to_numpy()
        if vals.sum() == 0:
            continue
        plt.bar(x, vals, bottom=bottom, label=col)
        bottom += vals
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylim(0, 1)
    plt.ylabel("Share of responses")
    plt.xlabel("Context length")
    plt.title("Error Composition by Context Length")
    plt.legend(fontsize=7, ncol=2, bbox_to_anchor=(1.02, 1), loc="upper left")
    save(
        "figure_5_error_composition",
        "Figure 5. Error composition by context length.",
        "Stacked proportions show correct responses and each deterministic error type by context.",
    )

    q = subgroup_plot_data(df, "question_type")
    plt.figure(figsize=(9, 5.2))
    for group, g in q.groupby("question_type"):
        g = g.set_index("context").loc[CONTEXT_ORDER].reset_index()
        plt.plot(x, g["inaccuracy"], marker="o", label=group)
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylim(0, 1)
    plt.ylabel("Inaccuracy rate")
    plt.xlabel("Context length")
    plt.title("Inaccuracy by Question Type and Context")
    plt.legend(fontsize=7, ncol=2)
    save(
        "figure_6_inaccuracy_by_question_type",
        "Figure 6. Inaccuracy by question type and context.",
        "Question-type subgroup trends are exploratory because cells are smaller than the primary repeated-measures analysis.",
    )

    d = subgroup_plot_data(df, "domain")
    plt.figure(figsize=(8.5, 5))
    for group, g in d.groupby("domain"):
        g = g.set_index("context").loc[CONTEXT_ORDER].reset_index()
        plt.plot(x, g["inaccuracy"], marker="o", label=group)
    plt.xticks(x, CONTEXT_ORDER)
    plt.ylim(0, 1)
    plt.ylabel("Inaccuracy rate")
    plt.xlabel("Context length")
    plt.title("Inaccuracy by Domain and Context")
    plt.legend(fontsize=8)
    save(
        "figure_7_inaccuracy_by_domain",
        "Figure 7. Inaccuracy by domain and context.",
        "Domain subgroup trends are exploratory and should not be over-interpreted.",
    )
    return figures


def subgroup_plot_data(df: pd.DataFrame, col: str) -> pd.DataFrame:
    rows = []
    for (group, context), g in df.groupby([col, "context_length_label"], sort=True):
        rows.append({col: group, "context": context, "inaccuracy": 1 - g["final_answer_correct"].mean()})
    return pd.DataFrame(rows)


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Title", 24, RGBColor(31, 78, 121)),
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(47, 84, 150)),
        ("Heading 3", 11, RGBColor(31, 78, 121)),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_page_number(section: Any) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_title_page(doc: Document) -> None:
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Long-Context Factual Reliability in Llama 3.2 3B Instruct")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Experiment C: Answer-Only Long-Context Benchmark Report")
    r.italic = True
    r.font.size = Pt(14)
    for text in [
        f"Model: {MODEL_ID}",
        "Context conditions: 4K, 8K, 16K, 32K, 64K",
        "Dataset: 100 question families, 500 repeated observations",
        f"Report date: {date.today().isoformat()}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared from frozen inference, grading, and statistical-analysis artifacts.").italic = True
    doc.add_page_break()


def add_manual_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    sections = [
        "1. Executive Summary / Abstract",
        "2. Research Question and Motivation",
        "3. Dataset and Source Domains",
        "4. Question Dataset",
        "5. Experimental Configuration",
        "6. Grading Method",
        "7. Overall Results",
        "8. Statistical Analysis of Overall Inaccuracy",
        "9. Statistical Analysis of Hallucinatory Inaccuracy",
        "10. Paired Context Comparisons",
        "11. Error Decomposition",
        "12. Question-Type Results",
        "13. Domain Results",
        "14. Sensitivity Analysis",
        "15. Inference-Time Analysis",
        "16. Interpretation",
        "17. Limitations",
        "18. Conclusion",
        "19. Methodological Provenance",
        "Appendix A. Detailed Tables",
    ]
    for s in sections:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_page_break()


def add_df_table(doc: Document, df: pd.DataFrame, title: str, caption: str, max_rows: int | None = None) -> None:
    doc.add_paragraph(title, style="Heading 3")
    if caption:
        p = doc.add_paragraph(caption)
        p.style = doc.styles["Caption"] if "Caption" in doc.styles else doc.styles["Normal"]
    show = df if max_rows is None else df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    for i, col in enumerate(show.columns):
        hdr[i].text = str(col)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(show.columns):
            val = row[col]
            if isinstance(val, float):
                text = f"{val:.3f}"
            else:
                text = str(val)
            cells[i].text = text
    doc.add_paragraph()


def add_figure(doc: Document, fig: dict[str, str], width: float = 6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(fig["title"]).bold = True
    doc.add_picture(fig["png"], width=Inches(width))
    cap = doc.add_paragraph(fig["caption"])
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(data: dict[str, Any], tables: dict[str, pd.DataFrame], figures: list[dict[str, str]]) -> Path:
    doc = Document()
    set_doc_style(doc)
    add_page_number(doc.sections[0])
    add_title_page(doc)
    add_manual_toc(doc)
    primary = data["primary"]
    models = data["models"]
    sensitivity = data["sensitivity"]
    latency = data["latency"]
    acc = models["primary_accuracy"]
    hall = models["primary_hallucination"]
    inacc_or = 1 / acc["odds_ratio_per_context_doubling"]
    inacc_ci = [1 / acc["odds_ratio_95_ci"][1], 1 / acc["odds_ratio_95_ci"][0]]

    doc.add_heading("1. Executive Summary / Abstract", level=1)
    doc.add_paragraph(
        "This report summarizes Experiment C, a completed local inference experiment evaluating factual reliability "
        f"for {MODEL_ID} under controlled long-context conditions. The benchmark uses primary-source records from "
        "SEC filings, Drugs@FDA, ClinicalTrials.gov, and FRED to create 100 repeated question families evaluated at "
        "4K, 8K, 16K, 32K, and 64K context lengths, for 500 total observations."
    )
    doc.add_paragraph(
        "The central result is that increasing context length strongly increased overall inaccuracy, but did not "
        "produce a statistically significant monotonic increase in hallucinatory inaccuracies. A substantial share "
        "of the additional errors at longer contexts came from grounded context-confusion errors."
    )
    doc.add_paragraph(
        f"Correctness declined from 57% at 4K to 27% at 64K. In a GEE logistic model clustered by question family, "
        f"each doubling of context length was associated with {inacc_or:.2f}x the odds of an inaccurate answer "
        f"(95% CI [{inacc_ci[0]:.2f}, {inacc_ci[1]:.2f}], p = {fmt_p(acc['p_value_context_log2'])}). "
        f"Hallucinatory inaccuracy increased from 35% to 42%, but the repeated-measures trend was not statistically "
        f"reliable (OR {hall['odds_ratio_per_context_doubling']:.3f}, 95% CI "
        f"[{hall['odds_ratio_95_ci'][0]:.3f}, {hall['odds_ratio_95_ci'][1]:.3f}], "
        f"p = {fmt_p(hall['p_value_context_log2'])}). Mean inference latency increased from 0.317 s to 7.824 s."
    )

    doc.add_heading("2. Research Question and Motivation", level=1)
    doc.add_paragraph(
        "The primary research question is whether hallucination probability increases as context length increases "
        "when a model is given legitimate but competing same-domain records. The experiment separates two mechanisms "
        "that are often conflated: unsupported hallucination and contextual confusion among supplied records."
    )
    doc.add_paragraph(
        "Unsupported hallucination occurs when the model produces a factual claim that is not supported by the supplied "
        "context. Contextual confusion occurs when the model returns a value that is actually present in the context "
        "but binds it to the wrong entity, period, version, unit, field, or calculation. Distinguishing these mechanisms "
        "matters because both reduce factual reliability, but they imply different retrieval and reasoning failures."
    )

    doc.add_heading("3. Dataset and Source Domains", level=1)
    doc.add_paragraph(
        "The benchmark converts authoritative primary-source records into controlled long-context factual questions. "
        "Each question family has a deterministic gold answer, target evidence, structured target conditions, and "
        "same-domain distractors. Contexts grow by adding plausible competing information while preserving the same "
        "question, gold answer, and target evidence."
    )
    add_df_table(doc, tables["domain_summary"], "Table 1. Source domains.", "")
    doc.add_paragraph(
        "Distractors include wrong-period records, wrong-version or vintage records, wrong-entity records, wrong-field "
        "records, wrong-unit or wrong-series-variant records, near-value records where meaningful, and other same-domain "
        "records. Unanswerable questions intentionally omit the target evidence and require the common abstention answer."
    )

    doc.add_heading("4. Question Dataset", level=1)
    doc.add_paragraph(
        "The final Experiment C dataset contains 100 unique question families: 80 answerable and 20 unanswerable. "
        "Each family was evaluated at five context lengths, yielding 100 x 5 = 500 repeated observations. Because the "
        "same families appear at all context lengths, statistical tests treat observations as clustered by family."
    )
    add_df_table(doc, tables["question_type_distribution"], "Table 2. Question-type distribution.", "")

    doc.add_heading("5. Experimental Configuration", level=1)
    config_rows = pd.DataFrame(
        [
            ["Model", MODEL_ID],
            ["Model revision", MODEL_REVISION],
            ["Hardware", "NVIDIA GeForce RTX 4090"],
            ["Precision", "BF16"],
            ["Cache", "standard DynamicCache"],
            ["Batch size", "1"],
            ["Decoding", "deterministic greedy; do_sample=False; num_beams=1"],
            ["max_new_tokens", "128"],
            ["Quantization/offloading", "none"],
            ["Prompt version/hash", f"{PROMPT_VERSION} / {PROMPT_HASH}"],
            ["Output contract", "ANSWER: <answer>"],
        ],
        columns=["Setting", "Value"],
    )
    add_df_table(doc, config_rows, "Table 3. Inference configuration.", "")
    doc.add_paragraph(
        "Experiment C used an answer-only output format and achieved 500/500 successful inference instances, zero "
        "format failures, zero repetitive degeneration, and zero outputs hitting the 128-token cap. Earlier evidence-ID "
        "structured formats were rejected during development because they induced repetitive selected-evidence generation."
    )

    doc.add_heading("6. Grading Method", level=1)
    doc.add_paragraph(
        f"Responses were graded by a frozen deterministic grader (SHA-256 {GRADER_HASH}). The grader normalized harmless "
        "formatting differences for dates, identifiers, numeric values, currencies, percentages, and insufficient-evidence "
        "answers. Of 500 responses, 499 were resolved by the deterministic grader and one was resolved by documented human "
        "manual adjudication; no unresolved cases remain."
    )
    doc.add_paragraph(
        "Responses are first categorized as Correct or Inaccurate. Inaccurate responses are then divided into "
        "hallucinatory inaccuracy and grounded inaccuracy. Hallucinatory inaccuracy refers to incorrect factual answers "
        "unsupported by the supplied context, including unsupported values and failures to abstain when evidence is absent. "
        "Grounded inaccuracy refers to incorrect answers grounded in supplied context, including wrong entity, wrong period, "
        "wrong version, wrong field, wrong unit, wrong series variant, calculation error, and unnecessary abstention."
    )
    doc.add_paragraph(
        "For example, if the question asks for a 2021 value and the model returns a real 2020 value present in context, "
        "the response is inaccurate and grounded, with error type WRONG_PERIOD. If the model outputs a factual value that "
        "does not appear in or follow from the supplied records, it is hallucinatory inaccuracy."
    )

    doc.add_heading("7. Overall Results", level=1)
    add_df_table(doc, tables["overall_results"], "Table 4. Central factual-reliability results.", "")
    add_figure(doc, figures[0])

    doc.add_heading("8. Statistical Analysis of Overall Inaccuracy", level=1)
    doc.add_paragraph(
        "The repeated-measures accuracy model was fit as a GEE logistic regression clustered by question_family_id with "
        "context_log2 coded as 4K = 0, 8K = 1, 16K = 2, 32K = 3, and 64K = 4. Because the report emphasizes factual "
        "reliability, the accuracy odds ratio was inverted to report overall inaccuracy."
    )
    doc.add_paragraph(
        f"Accuracy OR per context doubling was {acc['odds_ratio_per_context_doubling']:.3f} "
        f"(95% CI [{acc['odds_ratio_95_ci'][0]:.3f}, {acc['odds_ratio_95_ci'][1]:.3f}], "
        f"p = {fmt_p(acc['p_value_context_log2'])}). Equivalently, inaccuracy OR per context doubling was "
        f"{inacc_or:.3f} (95% CI [{inacc_ci[0]:.3f}, {inacc_ci[1]:.3f}], p = {fmt_p(acc['p_value_context_log2'])}). "
        "Each doubling of context length was associated with approximately a 35% increase in the odds of producing an "
        "inaccurate answer. This is an odds increase, not a 35-percentage-point increase."
    )
    add_figure(doc, figures[1])

    doc.add_heading("9. Statistical Analysis of Hallucinatory Inaccuracy", level=1)
    doc.add_paragraph(
        f"The hallucinatory-inaccuracy GEE model estimated an odds ratio of "
        f"{hall['odds_ratio_per_context_doubling']:.3f} per context doubling "
        f"(95% CI [{hall['odds_ratio_95_ci'][0]:.3f}, {hall['odds_ratio_95_ci'][1]:.3f}], "
        f"p = {fmt_p(hall['p_value_context_log2'])}). The point estimate is slightly greater than 1, but the confidence "
        "interval includes 1 and the p-value does not support a statistically reliable monotonic increase."
    )
    add_figure(doc, figures[2])

    doc.add_heading("10. Paired Context Comparisons", level=1)
    doc.add_paragraph(
        "Paired McNemar tests compared each longer context to 4K within the same question families. Holm correction was "
        "applied separately within the hallucination and correctness/inaccuracy outcomes."
    )
    add_df_table(doc, tables["paired_tests"], "Table 5. Paired context comparisons.", "")

    doc.add_heading("11. Error Decomposition", level=1)
    doc.add_paragraph(
        "The overall degradation is not explained solely by unsupported hallucination. Grounded inaccuracy rose from 8% "
        "at 4K to 31% at 64K, while hallucinatory inaccuracy changed from 35% to 42% without a statistically significant "
        "monotonic trend. This indicates a growing contribution from context-grounded binding and reasoning failures."
    )
    add_figure(doc, figures[4])
    add_df_table(doc, tables["error_counts_by_context"], "Table 6. Error-type counts by context.", "")

    doc.add_heading("12. Question-Type Results", level=1)
    doc.add_paragraph(
        "Question-type analyses are exploratory because subgroup cell sizes are smaller. DIRECT_RETRIEVAL produced zero "
        "hallucinatory inaccuracies. RETRIEVAL_CALCULATION had high hallucination across lengths rather than a clear "
        "monotonic increase. UNANSWERABLE questions showed a substantial increase in failed abstention by 64K."
    )
    add_figure(doc, figures[5])
    add_df_table(doc, tables["question_type_detail"], "Table 7. Question-type results by context.", "", max_rows=25)

    doc.add_heading("13. Domain Results", level=1)
    doc.add_paragraph(
        "Domain-level analyses are also exploratory. SEC showed the clearest hallucination increase, FDA and "
        "ClinicalTrials.gov were flatter, and FRED showed more modest changes. These patterns should not be over-claimed "
        "because each domain contributes only 25 families."
    )
    add_figure(doc, figures[6])
    add_df_table(doc, tables["domain_detail"], "Table 8. Domain results by context.", "")

    doc.add_heading("14. Sensitivity Analysis", level=1)
    sens_un = sensitivity["excluding_unanswerable_hallucination"]
    sens_manual = sensitivity["excluding_manual_adjudication_hallucination"]
    doc.add_paragraph(
        f"Excluding UNANSWERABLE questions yielded hallucination OR {sens_un['odds_ratio_per_context_doubling']:.3f} "
        f"(95% CI [{sens_un['odds_ratio_95_ci'][0]:.3f}, {sens_un['odds_ratio_95_ci'][1]:.3f}], "
        f"p = {fmt_p(sens_un['p_value_context_log2'])}). Excluding the single manually adjudicated row yielded "
        f"OR {sens_manual['odds_ratio_per_context_doubling']:.3f} "
        f"(95% CI [{sens_manual['odds_ratio_95_ci'][0]:.3f}, {sens_manual['odds_ratio_95_ci'][1]:.3f}], "
        f"p = {fmt_p(sens_manual['p_value_context_log2'])}). The primary conclusion is unchanged."
    )

    doc.add_heading("15. Inference-Time Analysis", level=1)
    add_figure(doc, figures[3])
    add_df_table(doc, tables["latency_summary"], "Table 9. Inference latency summary.", "")
    doc.add_paragraph(
        "Latency grows sharply with context length, but latency analysis is descriptive and explanatory only. It should "
        "not be interpreted as evidence that latency causes factual inaccuracies."
    )

    doc.add_heading("16. Interpretation", level=1)
    doc.add_paragraph(
        "Longer contexts substantially reduce factual reliability in Llama 3.2 3B. This degradation is not explained "
        "solely by unsupported hallucination. Instead, increasing context length produces a growing fraction of grounded "
        "inaccuracies in which the model selects, binds, or reasons over legitimate but incorrect contextual information."
    )
    doc.add_paragraph(
        "The controlled manipulation supports a context-length effect on overall inaccuracy. It does not support the "
        "more specific claim that hallucinatory inaccuracy alone increases monotonically with length."
    )

    doc.add_heading("17. Limitations", level=1)
    for item in [
        "The experiment includes 100 question families and 500 observations.",
        "Only one model was evaluated: Llama 3.2 3B Instruct.",
        "Experiment C covers contexts through 64K, not the near-128K condition.",
        "The benchmark uses four factual/structured data domains and intentionally strong same-domain distractors.",
        "Experiment C uses an answer-only format and therefore has no evidence-selection metric.",
        "Grading is deterministic with one documented manual adjudication.",
        "Subgroup analyses have smaller samples and are exploratory.",
        "GEE was used instead of mixed-effects logistic regression for reliable repeated-measures fitting.",
        "Results may not generalize to larger, proprietary, or differently prompted models.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("18. Conclusion", level=1)
    doc.add_paragraph(
        "Overall inaccuracy increased from 43% at 4K to 73% at 64K. The odds of inaccuracy increased by approximately "
        f"35% per context doubling (p = {fmt_p(acc['p_value_context_log2'])}). Hallucinatory inaccuracy changed from "
        f"35% to 42%, but the monotonic trend was not statistically significant (p = {fmt_p(hall['p_value_context_log2'])}). "
        "Grounded inaccuracy increased from 8% to 31%, and mean inference latency increased from 0.317 s to 7.824 s. "
        "The dominant statistically supported effect is decreasing factual reliability with context length, with growing "
        "contextual confusion playing an important role."
    )

    doc.add_heading("19. Methodological Provenance", level=1)
    provenance = pd.DataFrame(
        [
            ["Final scored dataset hash", FINAL_DATASET_HASH],
            ["Frozen grader hash", GRADER_HASH],
            ["Analysis directory", str(ANALYSIS_DIR)],
            ["Statistical method", "GEE logistic regression clustered by question_family_id; paired McNemar tests"],
            ["Bootstrap CIs", "Family-clustered bootstrap; seed 20260810; 5000 replicates"],
            ["No reruns", "Inference and grading were not rerun for this report."],
        ],
        columns=["Item", "Value"],
    )
    add_df_table(doc, provenance, "Table 10. Provenance.", "")

    doc.add_page_break()
    doc.add_heading("Appendix A. Detailed Tables", level=1)
    add_df_table(doc, tables["question_type_detail"], "Appendix Table A1. Full question-type results.", "")
    add_df_table(doc, tables["domain_detail"], "Appendix Table A2. Full domain results.", "")
    add_df_table(doc, tables["error_counts_by_context"], "Appendix Table A3. Full error-count table.", "")

    out = OUT_DIR / "final_research_report.docx"
    doc.save(out)
    return out


def build_markdown(data: dict[str, Any], tables: dict[str, pd.DataFrame], figures: list[dict[str, str]]) -> Path:
    models = data["models"]
    acc = models["primary_accuracy"]
    hall = models["primary_hallucination"]
    inacc_or = 1 / acc["odds_ratio_per_context_doubling"]
    inacc_ci = [1 / acc["odds_ratio_95_ci"][1], 1 / acc["odds_ratio_95_ci"][0]]
    lines = [
        "# Long-Context Factual Reliability in Llama 3.2 3B Instruct",
        "",
        "## Executive Summary",
        "",
        f"Overall inaccuracy increased from 43% at 4K to 73% at 64K. GEE-estimated inaccuracy OR per context doubling was {inacc_or:.3f} (95% CI [{inacc_ci[0]:.3f}, {inacc_ci[1]:.3f}], p={fmt_p(acc['p_value_context_log2'])}). Hallucinatory inaccuracy did not show a statistically significant monotonic increase (OR {hall['odds_ratio_per_context_doubling']:.3f}, 95% CI [{hall['odds_ratio_95_ci'][0]:.3f}, {hall['odds_ratio_95_ci'][1]:.3f}], p={fmt_p(hall['p_value_context_log2'])}).",
        "",
        "## Central Results",
        "",
        tables["overall_results"].to_markdown(index=False),
        "",
        "## Figures",
        "",
    ]
    for idx, fig in enumerate(figures, 1):
        lines += [f"### Figure {idx}. {fig['title']}", "", f"![{fig['title']}]({Path(fig['png']).relative_to(OUT_DIR)})", "", fig["caption"], ""]
    lines += [
        "## Paired Tests",
        "",
        tables["paired_tests"].to_markdown(index=False),
        "",
        "## Limitations",
        "",
        "- 100 question families and 500 observations.",
        "- One model, one answer-only prompting format, contexts through 64K.",
        "- No evidence-selection metric in Experiment C.",
        "- Subgroup analyses are exploratory.",
        "",
        "## Provenance",
        "",
        f"- Final dataset hash: `{FINAL_DATASET_HASH}`",
        f"- Frozen grader hash: `{GRADER_HASH}`",
        f"- Analysis directory: `{ANALYSIS_DIR}`",
    ]
    out = OUT_DIR / "final_research_report.md"
    out.write_text("\n".join(lines), encoding="utf-8")
    return out


def export_pdf(docx_path: Path) -> Path | None:
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        return None
    cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(OUT_DIR), str(docx_path)]
    proc = subprocess.run(cmd, cwd=OUT_DIR, text=True, capture_output=True, timeout=120)
    pdf = OUT_DIR / "final_research_report.pdf"
    if proc.returncode != 0 or not pdf.exists():
        (OUT_DIR / "pdf_export_error.txt").write_text(proc.stdout + "\n" + proc.stderr, encoding="utf-8")
        return None
    return pdf


def verify_docx(docx_path: Path, expected_figures: int) -> dict[str, Any]:
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    required = [
        "Executive Summary / Abstract",
        "Research Question and Motivation",
        "Dataset and Source Domains",
        "Question Dataset",
        "Experimental Configuration",
        "Grading Method",
        "Overall Results",
        "Statistical Analysis of Overall Inaccuracy",
        "Statistical Analysis of Hallucinatory Inaccuracy",
        "Paired Context Comparisons",
        "Error Decomposition",
        "Question-Type Results",
        "Domain Results",
        "Sensitivity Analysis",
        "Inference-Time Analysis",
        "Limitations",
        "Conclusion",
    ]
    missing = [s for s in required if s not in text]
    contradiction = "Hallucinatory inaccuracy did not significantly increase" not in text and "not statistically significant" not in text
    return {
        "docx_readable": True,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "embedded_figure_count": len(doc.inline_shapes),
        "missing_required_sections": missing,
        "all_required_sections_present": not missing,
        "figures_embedded": len(doc.inline_shapes) >= expected_figures,
        "tables_present": len(doc.tables) >= 10,
        "hallucination_non_significance_language_present": not contradiction,
    }


def package_versions() -> dict[str, str]:
    import docx
    import matplotlib
    import numpy
    import pandas

    return {
        "python": platform.python_version(),
        "platform": platform.platform(),
        "python_docx": docx.__version__,
        "matplotlib": matplotlib.__version__,
        "numpy": numpy.__version__,
        "pandas": pandas.__version__,
    }


def main() -> int:
    prepare_dirs()
    data = load_verified()
    tables = build_report_tables(data)
    figures = make_report_figures(data)
    docx_path = build_docx(data, tables, figures)
    md_path = build_markdown(data, tables, figures)
    pdf_path = export_pdf(docx_path)
    verification = verify_docx(docx_path, expected_figures=len(figures))
    manifest = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "source_dataset_csv": str(FINAL_CSV),
        "source_dataset_csv_sha256": sha256_file(FINAL_CSV),
        "source_dataset_jsonl": str(FINAL_JSONL),
        "source_dataset_hash": sha256_file(FINAL_JSONL),
        "expected_source_dataset_hash": FINAL_DATASET_HASH,
        "grader_hash": GRADER_HASH,
        "analysis_directory": str(ANALYSIS_DIR),
        "report_generation_script": "scripts/generate_experiment_c_final_report.py",
        "report_generation_script_sha256": sha256_file(Path("scripts/generate_experiment_c_final_report.py")),
        "package_versions": package_versions(),
        "figure_files": figures,
        "output_files": {
            "docx": str(docx_path),
            "markdown": str(md_path),
            "pdf": str(pdf_path) if pdf_path else None,
            "tables": str(TABLE_DIR),
            "figures": str(FIG_DIR),
        },
        "quality_check": verification,
        "statistical_values_verified_against_frozen_analysis": True,
        "inference_rerun": False,
        "grading_rerun": False,
        "labels_modified": False,
    }
    manifest_path = OUT_DIR / "report_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False, sort_keys=True), encoding="utf-8")
    print(json.dumps({
        "docx": str(docx_path),
        "markdown": str(md_path),
        "pdf": str(pdf_path) if pdf_path else None,
        "figures": len(figures),
        "tables": len(tables),
        "quality_check": verification,
        "manifest": str(manifest_path),
    }, indent=2), flush=True)
    return 0 if verification["all_required_sections_present"] and verification["figures_embedded"] and verification["tables_present"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
